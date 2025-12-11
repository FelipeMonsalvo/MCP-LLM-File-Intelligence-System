import io
import dropbox
from googleapiclient.http import MediaIoBaseDownload
from docx import Document

from drive_utils import (
    get_drive_service,
    get_first_5_folders_with_names as drive_get_first_5_folders_with_names,
    find_folder_by_name as drive_find_folder_by_name
)

from dropbox_utils import (
    get_dropbox_client,
    get_first_5_folders_with_names as dbx_get_first_5_folders_with_names,
    find_folder_by_name as dbx_find_folder_by_name
)


def list_files_fn(backend: str = "google", folder_id: str = None, folder_name: str = None):
    backend = backend.lower().strip()

    if backend == "google":
        service = get_drive_service()
        folder_ids = []

        if folder_name:
            target_folder_id = drive_find_folder_by_name(service, folder_name)
            if not target_folder_id:
                folders = drive_get_first_5_folders_with_names(service)
                names = [f["name"] for f in folders]
                return f"[Backend: Google Drive]\nFolder '{folder_name}' not found. Available: {', '.join(names)}"
            folder_ids = [target_folder_id]

        elif folder_id:
            folder_ids = [folder_id]

        else:
            folders = drive_get_first_5_folders_with_names(service)
            if not folders:
                return "[Backend: Google Drive]\nNo folders found in Google Drive."

            msg = "[Backend: Google Drive]\nFirst 5 Google Drive folders:\n\n"
            for f in folders:
                msg += f"- {f['name']} (ID: {f['id']})\n"
            return msg

        results_text = []
        for f_id in folder_ids[:5]:
            try:
                results = service.files().list(
                    q=f"'{f_id}' in parents and trashed=false",
                    pageSize=10,
                    fields="files(id, name, mimeType)"
                ).execute()

                files = results.get("files", [])
                results_text.append(f"\n[Backend: Google Drive]\nGoogle Folder: {f_id}\n")

                for f in files:
                    results_text.append(f"- {f['name']} (ID: {f['id']})")
            except Exception as e:
                results_text.append(f"[Backend: Google Drive]\nError: {e}")

        return "\n".join(results_text)

    elif backend == "dropbox":
        try:
            dbx = get_dropbox_client()
        except RuntimeError as e:
            error_msg = str(e)
            if "DROPBOX_ACCESS_TOKEN is missing" in error_msg:
                return "[Backend: Dropbox]\nDropbox access not configured. Add DROPBOX_ACCESS_TOKEN to mcp_server/.env"
            return f"[Backend: Dropbox]\nDropbox authentication error: {error_msg}"
        except Exception as e:
            return f"[Backend: Dropbox]\nError connecting to Dropbox: {str(e)}"

        target_path = None

        if folder_name:
            target_path = dbx_find_folder_by_name(dbx, folder_name)
            if not target_path:
                folders = dbx_get_first_5_folders_with_names(dbx)
                names = [f["name"] for f in folders]
                return f"[Backend: Dropbox]\nDropbox folder '{folder_name}' not found. Available: {', '.join(names)}"

        elif folder_id:
            target_path = folder_id.strip()
            if not target_path.startswith("/"):
                target_path = "/" + target_path
            target_path = target_path.lower()

        if not target_path:
            try:
                result = dbx.files_list_folder("", recursive=False)

                folders = []
                files = []

                for entry in result.entries:
                    if isinstance(entry, dropbox.files.FolderMetadata):
                        folders.append(entry)
                    elif isinstance(entry, dropbox.files.FileMetadata):
                        files.append(entry)

                msg = "[Backend: Dropbox]\nDropbox Root Contents:\n\n"

                if folders:
                    msg += f"Folders (showing first {min(len(folders), 10)}):\n"
                    for f in folders[:10]:
                        msg += f"- {f.name} (Use folder_id: '{f.path_lower}' to open)\n"
                    if len(folders) > 10:
                        msg += f"\n... and {len(folders) - 10} more folders\n"

                if files:
                    msg += "\nFiles:\n"
                    for f in files[:10]:
                        msg += f"- {f.name}\n"
                    if len(files) > 10:
                        msg += f"\n... and {len(files) - 10} more files\n"

                if not folders and not files:
                    msg += "Root folder is empty.\n"

                return msg

            except Exception as e:
                return f"[Backend: Dropbox]\nError listing Dropbox root: {e}"

        try:
            result = dbx.files_list_folder(target_path, recursive=False)

            folders = []
            files = []

            for entry in result.entries:
                if isinstance(entry, dropbox.files.FolderMetadata):
                    folders.append(entry)
                elif isinstance(entry, dropbox.files.FileMetadata):
                    files.append(entry)

            msg = f"[Backend: Dropbox]\nDropbox Folder: {target_path}\n\n"

            if folders:
                msg += "Folders:\n"
                for f in folders[:10]:
                    msg += f"- {f.name} (Use folder_id: '{f.path_lower}' to open)\n"

            if files:
                msg += "\nFiles:\n"
                for f in files[:10]:
                    msg += f"- {f.name}\n"

            if not folders and not files:
                msg += "This folder is empty.\n"

            return msg

        except Exception as e:
            error_str = str(e)
            if "not_found" in error_str.lower() or "not found" in error_str.lower():
                return f"[Backend: Dropbox]\nError: Path '{target_path}' not found or inaccessible."
            return f"[Backend: Dropbox]\nError accessing Dropbox folder: {error_str}"

    return "Invalid backend."


def search_files_fn(backend: str = "google", query: str = "", folder_id: str = None, folder_name: str = None):
    backend = backend.lower().strip()
    query = query.lower().strip()

    if backend == "google":
        service = get_drive_service()
        folder_ids = []

        if folder_name:
            target_folder_id = drive_find_folder_by_name(service, folder_name)
            if not target_folder_id:
                folders = drive_get_first_5_folders_with_names(service)
                names = [f["name"] for f in folders]
                return f"[Backend: Google Drive]\nFolder '{folder_name}' not found. Available: {', '.join(names)}"
            folder_ids = [target_folder_id]

        elif folder_id:
            folder_ids = [folder_id]

        else:
            folders = drive_get_first_5_folders_with_names(service)

            if not folders:
                return "[Backend: Google Drive]\nNo folders found in Google Drive."

            msg = "[Backend: Google Drive]\nFirst 5 Google Drive folders:\n\n"

            for f in folders:
                msg += f"- {f['name']} (ID: {f['id']})\n"

            msg += "\nPlease specify a folder to search."
            return msg

        results_text = []

        for f_id in folder_ids[:5]:
            full_query = (
                f"'{f_id}' in parents and "
                f"(name contains '{query}' or fullText contains '{query}') and trashed=false"
            )

            try:
                results = service.files().list(
                    q=full_query,
                    pageSize=10,
                    fields="files(id, name, mimeType, modifiedTime)"
                ).execute()

                files = results.get("files", [])
                if not files:
                    continue

                results_text.append(f"\n[Backend: Google Drive]\nGoogle Folder: {f_id}\n")
                for f in files:
                    results_text.append(
                        f"- {f['name']} (ID: {f['id']}, Type: {f['mimeType']})"
                    )

            except Exception as e:
                results_text.append(f"[Backend: Google Drive]\nError searching folder {f_id}: {e}")

        return "\n".join(results_text) if results_text else "[Backend: Google Drive]\nNo matching files found."

    elif backend == "dropbox":
        try:
            dbx = get_dropbox_client()
        except RuntimeError as e:
            error_msg = str(e)
            if "DROPBOX_ACCESS_TOKEN is missing" in error_msg:
                return "[Backend: Dropbox]\nDropbox access not configured. Add DROPBOX_ACCESS_TOKEN to mcp_server/.env"
            return f"[Backend: Dropbox]\nDropbox authentication error: {error_msg}"
        except Exception as e:
            return f"[Backend: Dropbox]\nError connecting to Dropbox: {str(e)}"

        target_path = None

        if folder_name:
            target_path = dbx_find_folder_by_name(dbx, folder_name)
            if not target_path:
                try:
                    result = dbx.files_list_folder("", recursive=False)
                    folders = [e for e in result.entries if isinstance(e, dropbox.files.FolderMetadata)]
                    names = [f.name for f in folders]
                except:
                    names = []
                return f"[Backend: Dropbox]\nDropbox folder '{folder_name}' not found. Available: {', '.join(names)}"

        elif folder_id:
            target_path = folder_id.strip()
            
            if not target_path.startswith("/"):
                target_path = "/" + target_path
            target_path = target_path.lower()

        if not target_path:
            try:
                result = dbx.files_list_folder("", recursive=False)
                files = [e for e in result.entries if isinstance(e, dropbox.files.FileMetadata)]

                matched = [f for f in files if query in f.name.lower()]

                if matched:
                    msg = "[Backend: Dropbox]\nDropbox Root Search Results:\n\n"
                    for f in matched:
                        msg += f"- {f.name} (Path: {f.path_lower})\n"
                    return msg

                return f"[Backend: Dropbox]\nNo root Dropbox files match '{query}'."

            except Exception as e:
                return f"[Backend: Dropbox]\nError searching Dropbox root: {e}"

        try:
            result = dbx.files_list_folder(target_path, recursive=False)
            files = [e for e in result.entries if isinstance(e, dropbox.files.FileMetadata)]

            matched = [f for f in files if query in f.name.lower()]

            if not matched:
                return f"[Backend: Dropbox]\nNo Dropbox files in '{target_path}' match '{query}'."

            msg = f"[Backend: Dropbox]\nDropbox Folder Search: {target_path}\n\n"
            for f in matched:
                msg += f"- {f.name} (Path: {f.path_lower})\n"

            return msg

        except Exception as e:
            error_str = str(e)
            if "not_found" in error_str.lower() or "not found" in error_str.lower():
                return f"[Backend: Dropbox]\nError: Path '{target_path}' not found or inaccessible."
            return f"[Backend: Dropbox]\nError searching Dropbox folder: {error_str}"


def drive_read_file(service, file_id):
    meta = service.files().get(fileId=file_id, fields="id,name,mimeType").execute()
    mime = meta["mimeType"]
    name = meta["name"]

    fh = io.BytesIO()

    if mime == "application/vnd.google-apps.document":
        request = service.files().export_media(fileId=file_id, mimeType="text/plain")
        downloader = MediaIoBaseDownload(fh, request)
    elif mime == "text/plain" or name.lower().endswith(".md"):
        request = service.files().get_media(fileId=file_id)
        downloader = MediaIoBaseDownload(fh, request)
    elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or name.lower().endswith(".docx"):
        request = service.files().get_media(fileId=file_id)
        downloader = MediaIoBaseDownload(fh, request)
    else:
        return None, f"Unsupported Google Drive file type: {mime}"

    done = False
    while not done:
        _, done = downloader.next_chunk()
    fh.seek(0)

    if name.lower().endswith(".docx"):
        doc = Document(fh)
        text = "\n".join(p.text for p in doc.paragraphs)
        return text, name

    text = fh.read().decode("utf-8", errors="replace")
    return text, name


def dropbox_read_file(dbx, file_path):
    if not file_path:
        return None, "file_path is required for Dropbox files"

    normalized_path = file_path.strip()
    if not normalized_path.startswith("/"):
        normalized_path = "/" + normalized_path

    try:
        _, result = dbx.files_download(normalized_path)
        raw = result.content
        name = normalized_path.split("/")[-1]

        if name.lower().endswith((".txt", ".md")):
            return raw.decode("utf-8", errors="replace"), name

        if name.lower().endswith(".docx"):
            doc = Document(io.BytesIO(raw))
            text = "\n".join(p.text for p in doc.paragraphs)
            return text, name

        return None, f"Unsupported Dropbox file type: {name}"

    except Exception as e:
        return None, f"Error reading Dropbox file '{normalized_path}': {e}"


def get_file_fn(
    backend: str,
    file_id: str = None,
    file_path: str = None
):
    backend = backend.lower().strip()

    if backend == "google":
        service = get_drive_service()
        text, name = drive_read_file(service, file_id)

        if text is None:
            return name
        return f"[Google Drive File: {name}]\n\n{text}"

    elif backend == "dropbox":
        dbx = get_dropbox_client()
        text, name = dropbox_read_file(dbx, file_path)

        if text is None:
            return name
        return f"[Dropbox File: {name}]\n\n{text}"

    else:
        return "Invalid backend. Use 'google' or 'dropbox'."


def summarize_file_fn(
    backend: str,
    file_id: str = None,
    file_path: str = None
) -> str:
    text = ""
    file_name = ""

    if backend == "google":
        try:
            service = get_drive_service()

            meta = service.files().get(
                fileId=file_id,
                fields="id, name, mimeType"
            ).execute()

            file_name = meta["name"]
            mime = meta["mimeType"]

            if mime == "text/plain" or file_name.lower().endswith(".txt"):
                request = service.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                fh.seek(0)
                text = fh.read().decode("utf-8", errors="ignore")

            elif file_name.lower().endswith(".docx"):
                request = service.files().get_media(fileId=file_id)
                fh = io.BytesIO()
                downloader = MediaIoBaseDownload(fh, request)
                done = False
                while not done:
                    _, done = downloader.next_chunk()
                fh.seek(0)
                doc = Document(fh)
                text = "\n".join(p.text for p in doc.paragraphs)

            else:
                return f"Unsupported file type: {file_name}. Only .txt and .docx are supported."
        except Exception as e:
            error_str = str(e)
            if "not found" in error_str.lower() or "404" in error_str.lower():
                return f"[Backend: Google Drive]\nError: File with ID '{file_id}' not found or inaccessible."
            return f"[Backend: Google Drive]\nError reading Google Drive file: {error_str}"

    elif backend == "dropbox":
        try:
            dbx = get_dropbox_client()
        except RuntimeError as e:
            error_msg = str(e)
            if "DROPBOX_ACCESS_TOKEN is missing" in error_msg:
                return "[Backend: Dropbox]\nDropbox access not configured. Add DROPBOX_ACCESS_TOKEN to mcp_server/.env"
            return f"[Backend: Dropbox]\nDropbox authentication error: {error_msg}"
        except Exception as e:
            return f"[Backend: Dropbox]\nError connecting to Dropbox: {str(e)}"

        if not file_path:
            return "file_path is required for Dropbox files"

        normalized_path = file_path.strip()
        if not normalized_path.startswith("/"):
            normalized_path = "/" + normalized_path

        try:
            md, response = dbx.files_download(normalized_path)
            file_name = normalized_path.split("/")[-1]
            raw = response.content

            if file_name.lower().endswith(".txt"):
                text = raw.decode("utf-8", errors="ignore")

            elif file_name.lower().endswith(".docx"):
                buf = io.BytesIO(raw)
                doc = Document(buf)
                text = "\n".join(p.text for p in doc.paragraphs)

            else:
                return f"Unsupported file type: {file_name}. Only .txt and .docx are supported."
        except Exception as e:
            error_str = str(e)
            if "not_found" in error_str.lower() or "not found" in error_str.lower():
                return f"[Backend: Dropbox]\nError: File path '{normalized_path}' not found or inaccessible."
            return f"[Backend: Dropbox]\nError reading Dropbox file '{normalized_path}': {error_str}"

    else:
        return "Invalid backend."

    if not text.strip():
        return f"The file '{file_name}' contains no readable text."

    return (
        f"File: {file_name}\n"
        f"Backend: {backend}\n"
        f"Content:\n\n{text}"
    )
